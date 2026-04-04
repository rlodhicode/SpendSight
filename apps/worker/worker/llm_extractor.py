from __future__ import annotations
import json
from abc import ABC, abstractmethod
from datetime import date
from pathlib import Path
from tempfile import NamedTemporaryFile
from typing import Literal, Optional

from google import genai
from google.genai import types
from pydantic import BaseModel, Field, ValidationError

from .config import settings


class BillingPeriod(BaseModel):
    start_date: Optional[date] = None
    end_date: Optional[date] = None


class Header(BaseModel):
    utility_type: Optional[Literal["electric", "water", "gas", "waste", "internet"]] = None
    provider_name: Optional[str] = None
    account_number: Optional[str] = None
    billing_period: BillingPeriod = Field(default_factory=BillingPeriod)
    due_date: Optional[date] = None


class MeterReading(BaseModel):
    meter_id: Optional[str] = None
    start: Optional[float] = None
    end: Optional[float] = None


class Usage(BaseModel):
    amount: Optional[float] = None
    unit: Optional[Literal["kWh", "Gallons", "Therms"]] = None
    meter_readings: list[MeterReading] = Field(default_factory=list)


class Adjustment(BaseModel):
    description: Optional[str] = None
    amount: Optional[float] = None


class Financials(BaseModel):
    previous_balance: Optional[float] = None
    payments_credits: Optional[float] = None
    current_charges: Optional[float] = None
    total_amount_due: Optional[float] = None
    adjustments: list[Adjustment] = Field(default_factory=list)


class LineItem(BaseModel):
    description: Optional[str] = None
    quantity: Optional[float] = None
    rate: Optional[float] = None
    total: Optional[float] = None


class UtilityBillSchema(BaseModel):
    header: Header
    usage: Usage
    financials: Financials
    line_items: list[LineItem] = Field(default_factory=list)


class FieldConfidence(BaseModel):
    field_name: str
    confidence_score: float
    source: str = "gemini"


class ExtractionEnvelope(BaseModel):
    extracted: UtilityBillSchema
    field_confidences: list[FieldConfidence] = Field(default_factory=list)


class ExtractedBillDocument(BaseModel):
    extracted: UtilityBillSchema
    field_confidences: list[FieldConfidence] = Field(default_factory=list)
    overall_confidence: float = 0.0


class ExtractionProvider(ABC):
    @abstractmethod
    def extract(self, data: bytes, filename: str) -> ExtractedBillDocument:
        raise NotImplementedError


def _mime_for_filename(filename: str) -> str:
    ext = Path(filename).suffix.lower()
    mapping = {
        ".pdf": "application/pdf",
        ".png": "image/png",
        ".jpg": "image/jpeg",
        ".jpeg": "image/jpeg",
        ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    }
    return mapping.get(ext, "application/octet-stream")


def _docx_to_text(data: bytes) -> str:
    try:
        from docx import Document as DocxDocument
    except Exception as exc:
        raise RuntimeError("DOCX parsing requires python-docx/lxml runtime support") from exc
    with NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        tmp.write(data)
        tmp_path = Path(tmp.name)
    try:
        doc = DocxDocument(str(tmp_path))
        return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
    finally:
        if tmp_path.exists():
            tmp_path.unlink()


def _prompt_schema() -> str:
    schema = {
        "extracted": {
            "header": {
                "utility_type": "electric | water | gas | waste | internet",
                "provider_name": "string",
                "account_number": "string",
                "billing_period": {"start_date": "YYYY-MM-DD", "end_date": "YYYY-MM-DD"},
                "due_date": "YYYY-MM-DD",
            },
            "usage": {
                "amount": 0.0,
                "unit": "kWh | Gallons | Therms",
                "meter_readings": [{"meter_id": "string", "start": 0.0, "end": 0.0}],
            },
            "financials": {
                "previous_balance": 0.0,
                "payments_credits": 0.0,
                "current_charges": 0.0,
                "total_amount_due": 0.0,
                "adjustments": [{"description": "string", "amount": 0.0}],
            },
            "line_items": [{"description": "string", "quantity": 0.0, "rate": 0.0, "total": 0.0}],
        },
        "field_confidences": [
            {"field_name": "header.provider_name", "confidence_score": 0.0, "source": "gemini"},
            {"field_name": "header.billing_period.start_date", "confidence_score": 0.0, "source": "gemini"},
            {"field_name": "header.billing_period.end_date", "confidence_score": 0.0, "source": "gemini"},
            {"field_name": "financials.total_amount_due", "confidence_score": 0.0, "source": "gemini"},
            {"field_name": "usage.amount", "confidence_score": 0.0, "source": "gemini"},
        ],
    }
    return (
        "You are an expert utility bill auditor.\n"
        "Extract data from the provided bill document.\n"
        "Return strictly valid JSON matching this exact schema.\n"
        f"{json.dumps(schema)}\n"
        "Return confidence_score values from 0.0 to 1.0.\n"
        "If a scalar field is missing, return null.\n"
        "If a list field is missing, return an empty array [].\n"
        "Do not return markdown."
    )


def _normalize_payload(parsed: dict) -> dict:
    envelope = parsed
    if "extracted" not in envelope or envelope["extracted"] is None:
        envelope["extracted"] = {}
    extracted = envelope["extracted"]

    if extracted.get("header") is None:
        extracted["header"] = {}
    if extracted["header"].get("billing_period") is None:
        extracted["header"]["billing_period"] = {}

    if extracted.get("usage") is None:
        extracted["usage"] = {}
    if extracted["usage"].get("meter_readings") is None:
        extracted["usage"]["meter_readings"] = []

    if extracted.get("financials") is None:
        extracted["financials"] = {}
    if extracted["financials"].get("adjustments") is None:
        extracted["financials"]["adjustments"] = []

    if extracted.get("line_items") is None:
        extracted["line_items"] = []

    # Normalize enum-like values to avoid hard-failing extraction on blurry/OCR-noisy docs.
    utility_type = extracted["header"].get("utility_type")
    if utility_type not in {"electric", "water", "gas", "waste", "internet", None}:
        extracted["header"]["utility_type"] = None

    usage_unit = extracted["usage"].get("unit")
    if usage_unit not in {"kWh", "Gallons", "Therms", None}:
        extracted["usage"]["unit"] = None

    if envelope.get("field_confidences") is None:
        envelope["field_confidences"] = []

    return envelope


def _empty_extracted_document() -> ExtractedBillDocument:
    extracted = UtilityBillSchema.model_validate(
        {
            "header": {},
            "usage": {"meter_readings": []},
            "financials": {"adjustments": []},
            "line_items": [],
        }
    )
    confidences = [
        FieldConfidence(field_name="header.provider_name", confidence_score=0.05, source="validation-fallback"),
        FieldConfidence(
            field_name="header.billing_period.start_date",
            confidence_score=0.05,
            source="validation-fallback",
        ),
        FieldConfidence(
            field_name="header.billing_period.end_date",
            confidence_score=0.05,
            source="validation-fallback",
        ),
        FieldConfidence(field_name="financials.total_amount_due", confidence_score=0.05, source="validation-fallback"),
        FieldConfidence(field_name="usage.amount", confidence_score=0.05, source="validation-fallback"),
        FieldConfidence(field_name="usage.unit", confidence_score=0.05, source="validation-fallback"),
    ]
    return ExtractedBillDocument(extracted=extracted, field_confidences=confidences, overall_confidence=0.05)


def _fallback_field_confidences(extracted: UtilityBillSchema) -> list[FieldConfidence]:
    # Fallback confidence map for non-conforming model output.
    required_pairs = [
        ("header.provider_name", extracted.header.provider_name),
        ("header.account_number", extracted.header.account_number),
        ("header.billing_period.start_date", extracted.header.billing_period.start_date),
        ("header.billing_period.end_date", extracted.header.billing_period.end_date),
        ("financials.total_amount_due", extracted.financials.total_amount_due),
        ("usage.amount", extracted.usage.amount),
        ("usage.unit", extracted.usage.unit),
    ]
    rows: list[FieldConfidence] = []
    for field_name, value in required_pairs:
        rows.append(
            FieldConfidence(
                field_name=field_name,
                confidence_score=0.6 if value is not None else 0.2,
                source="gemini-fallback",
            )
        )
    return rows


def _compute_overall_confidence(confidences: list[FieldConfidence]) -> float:
    if not confidences:
        return 0.0
    total = sum(max(0.0, min(1.0, item.confidence_score)) for item in confidences)
    return round(total / len(confidences), 4)


class VertexGeminiExtractor(ExtractionProvider):
    def extract(self, data: bytes, filename: str) -> ExtractedBillDocument:
        client = genai.Client(vertexai=True, project=settings.gcp_project_id, location=settings.gcp_location)
        mime = _mime_for_filename(filename)
        prompt = _prompt_schema()

        if mime.endswith("document.wordprocessingml.document"):
            contents = [f"Document text:\n{_docx_to_text(data)}\n\n{prompt}"]
        else:
            contents = [types.Part.from_bytes(data=data, mime_type=mime), prompt]

        response = client.models.generate_content(
            model=settings.gemini_model,
            contents=contents,
            config=types.GenerateContentConfig(response_mime_type="application/json"),
        )
        raw_json = response.text or "{}"
        try:
            parsed = json.loads(raw_json)
        except json.JSONDecodeError as exc:
            raise ValueError(f"Model returned invalid JSON: {exc}") from exc
        if not isinstance(parsed, dict):
            raise ValueError("Model output must be a JSON object")
        parsed = _normalize_payload(parsed)

        try:
            envelope = ExtractionEnvelope.model_validate(parsed)
        except ValidationError as exc:
            # Graceful degradation: return empty/low-confidence extraction so flow routes to human review.
            return _empty_extracted_document()

        field_confidences = envelope.field_confidences or _fallback_field_confidences(envelope.extracted)
        return ExtractedBillDocument(
            extracted=envelope.extracted,
            field_confidences=field_confidences,
            overall_confidence=_compute_overall_confidence(field_confidences),
        )


